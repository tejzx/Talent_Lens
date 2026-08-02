"""DOCX text extraction (paragraphs + tables)."""
from __future__ import annotations

import io


def extract_docx_text(data: bytes) -> str:
    from docx import Document  # imported lazily so the package stays optional

    document = Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()
